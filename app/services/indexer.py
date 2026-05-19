"""FAISS index builder.

Pipeline:
  1. Read every active file from settings.uploads_path.
  2. Extract text:  .docx via python-docx, .csv via pandas.
  3. Chunk into ~500-word windows (matches the original v1 indexer).
  4. Embed via the shared sentence-transformers model (same one used at query time).
  5. Build IndexFlatL2.
  6. Write `<timestamp>__trip_safe_index.faiss` + `.pkl` to faiss_index_dir as a
     temp pair, then rename to the canonical names so the active index can't be
     left in a half-written state.
  7. Call rag_engine.reload() — the next chat picks up the new index without a
     restart. The previous .faiss/.pkl are kept under faiss_index_dir/backups/.

The whole job runs in a FastAPI BackgroundTask. CRITICAL: every blocking call
(model load, embedding, FAISS build, disk IO) is wrapped in asyncio.to_thread
so the event loop stays free — otherwise the entire backend freezes while a
rebuild is in progress.
"""
from __future__ import annotations

import asyncio
import logging
import pickle
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import numpy as np
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.models.index_build import IndexBuild, IndexBuildStatus
from app.models.knowledge_file import KnowledgeFile
from app.services.rag import (
    INDEX_FILENAME,
    METADATA_FILENAME,
    engine as rag_engine,
)

logger = logging.getLogger("tripsafe.indexer")

CHUNK_WORD_SIZE = 500
CHUNK_OVERLAP_WORDS = 50


# ── Text extraction ─────────────────────────────────────────────────────
#
# Two-mode strategy: PROSE paragraphs go into a buffered text stream that's
# sliced into ~500-word windows like before. TABLE rows are emitted as their
# own structured "one row per chunk" chunks, with the table caption (the
# paragraph immediately preceding the table) and the column headers baked
# into the chunk text. This is the fix for the "pricing is wrong" class of
# bug — each price now lives in a self-contained chunk like:
#
#   Document: 03 - TripSafe Pricing Info (Main Plans).docx
#   Table: Classic Silver ($50,000) — Region: Asia
#   Trip Duration: 1-4 | 0-40 yrs: ₹470 | 41-60 yrs: ₹520 | 61-70 yrs: ₹750 | 71-75 yrs: ₹1,390
#
# which is trivially retrievable by both embedding similarity and keyword match.


def _para_chunks_from_buffer(
    buf: list[str], source: str, section_heading: str
) -> list[tuple[str, str]]:
    """Slice an accumulated prose buffer into 500-word windows.
    Every emitted chunk is prefixed with `Document:` + `Section:` so the
    LLM always knows where the text came from."""
    text = "\n".join(buf).strip()
    if not text:
        return []
    words = text.split()
    out: list[tuple[str, str]] = []
    step = CHUNK_WORD_SIZE - CHUNK_OVERLAP_WORDS
    header_lines = [f"Document: {source}"]
    if section_heading:
        header_lines.append(f"Section: {section_heading}")
    header = "\n".join(header_lines) + "\n\n"
    for i in range(0, len(words), step):
        window = words[i : i + CHUNK_WORD_SIZE]
        if not window:
            break
        out.append((header + " ".join(window), source))
        if i + CHUNK_WORD_SIZE >= len(words):
            break
    return out


def _extract_docx_chunks(path: Path, source: str) -> list[tuple[str, str]]:
    """Walk the docx body in document order. Paragraphs accumulate into a
    buffer that's flushed (a) when we hit a table, or (b) at end of doc.
    Tables emit one chunk per data row, with the immediately-preceding
    paragraph used as the table caption."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(path))

    chunks: list[tuple[str, str]] = []
    para_buffer: list[str] = []
    last_para_text: str = ""  # candidate table caption
    section_heading: str = ""  # heading-styled paragraph (broader context)

    def flush() -> None:
        nonlocal para_buffer
        if para_buffer:
            chunks.extend(
                _para_chunks_from_buffer(para_buffer, source, section_heading)
            )
            para_buffer = []

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name or "title" in style_name:
                # Promote to section heading. Flush prior prose so it's
                # filed under the old heading.
                flush()
                section_heading = text
            para_buffer.append(text)
            last_para_text = text
        elif child.tag == qn("w:tbl"):
            # Flush prose first so it's not mashed with table rows.
            flush()
            table = Table(child, doc)
            if not table.rows:
                last_para_text = ""
                continue
            # First row is the column header. Some docx tables don't have a
            # proper header row — if every cell in row 0 is empty, fall back
            # to generic "Col 1, Col 2…" labels so we still emit row data.
            header_cells = [c.text.strip() for c in table.rows[0].cells]
            if not any(header_cells):
                header_cells = [
                    f"Col {i+1}" for i in range(len(table.rows[0].cells))
                ]
            data_rows = table.rows[1:] if any(
                c.text.strip() for c in table.rows[0].cells
            ) else table.rows
            caption = last_para_text or section_heading
            for row in data_rows:
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                pairs: list[str] = []
                for h, v in zip(header_cells, cells):
                    if not v:
                        continue
                    pairs.append(f"{h}: {v}" if h else v)
                if not pairs:
                    continue
                lines = [f"Document: {source}"]
                if section_heading and section_heading != caption:
                    lines.append(f"Section: {section_heading}")
                if caption:
                    lines.append(f"Table: {caption}")
                lines.append("Row: " + " | ".join(pairs))
                chunks.append(("\n".join(lines), source))
            last_para_text = ""

    flush()
    return chunks


def _extract_csv_chunks(path: Path, source: str) -> list[tuple[str, str]]:
    """One chunk per row, with column headers labelling each cell value.
    Same logic as docx tables — keeps each fact retrievable in isolation."""
    import pandas as pd

    df = pd.read_csv(path, dtype=str, keep_default_na=False)
    headers = [str(c) for c in df.columns]
    chunks: list[tuple[str, str]] = []
    for _, row in df.iterrows():
        pairs: list[str] = []
        for h, v in zip(headers, row.values):
            v_str = str(v).strip()
            if v_str:
                pairs.append(f"{h}: {v_str}")
        if not pairs:
            continue
        text = f"Document: {source}\nRow: " + " | ".join(pairs)
        chunks.append((text, source))
    return chunks


def extract_chunks(path: Path, source: str) -> list[tuple[str, str]]:
    """Returns ready-to-embed (chunk_text, source) tuples for a single file."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx_chunks(path, source)
    if suffix == ".csv":
        return _extract_csv_chunks(path, source)
    raise ValueError(f"Unsupported file type: {suffix}")


# ── Legacy helpers (kept for any external callers / tests) ──────────────

def _extract_docx(path: Path) -> str:
    """Legacy flat-text extractor. The real ingestion path now uses
    extract_chunks() which is structure-aware. This exists only so external
    callers (e.g. ad-hoc scripts) don't break."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    import pandas as pd

    df = pd.read_csv(path, dtype=str, keep_default_na=False)
    lines: list[str] = [" | ".join(str(c) for c in df.columns)]
    for _, row in df.iterrows():
        lines.append(" | ".join(str(v) for v in row.values))
    return "\n".join(lines)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".csv":
        return _extract_csv(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def chunk_text(text: str, source: str) -> list[tuple[str, str]]:
    """Legacy. Only used if a caller still hands in raw text."""
    words = text.split()
    if not words:
        return []
    chunks: list[tuple[str, str]] = []
    step = CHUNK_WORD_SIZE - CHUNK_OVERLAP_WORDS
    for i in range(0, len(words), step):
        window = words[i : i + CHUNK_WORD_SIZE]
        if not window:
            break
        chunks.append((" ".join(window), source))
        if i + CHUNK_WORD_SIZE >= len(words):
            break
    return chunks


# ── Main build routine ──────────────────────────────────────────────────

async def run_build(
    db: AsyncSession,
    build_id: uuid.UUID,
) -> None:
    """Long-running task. Updates the IndexBuild row as it progresses so the
    admin panel can poll status. Never raises — always settles the row to
    COMPLETE or FAILED."""
    settings = get_settings()
    build = await _get_build(db, build_id)
    if build is None:
        logger.error("index_build_missing id=%s", build_id)
        return
    build.status = IndexBuildStatus.RUNNING
    await db.commit()

    try:
        # 1. List active source files in DB.
        files_result = await db.execute(
            select(KnowledgeFile).where(KnowledgeFile.is_active.is_(True))
        )
        files = list(files_result.scalars().all())
        if not files:
            raise RuntimeError("No active source files to index.")

        # 2. Extract + chunk. extract_chunks() is table-aware: docx/csv tables
        # produce one chunk per row with column headers baked in, while prose
        # paragraphs still get the 500-word window treatment. This is what
        # makes pricing-style "lookup" questions actually work — each row's
        # numbers stay tied to their plan/region/age band labels.
        all_chunks: list[tuple[str, str]] = []
        source_summary: list[dict[str, Any]] = []
        from app.services import s3_store

        for f in files:
            path = Path(f.stored_path)
            # Container restarts wipe /app/uploads. If S3 sync is enabled,
            # transparently pull the file back from S3 before extracting.
            if not path.exists():
                s3_store.ensure_upload_local(path)
            else:
                # File exists locally but may not yet be backed up to S3
                # (e.g. it was uploaded before the S3-uploads feature shipped).
                # push_upload is a cheap no-op if S3 uploads is disabled, and
                # boto3.upload_file is content-stable so re-pushes are harmless.
                s3_store.push_upload(path)
            if not path.exists():
                logger.warning("source_file_missing path=%s", path)
                continue
            try:
                file_chunks = extract_chunks(path, source=f.filename)
            except Exception:
                logger.exception("extract_failed file=%s", f.filename)
                continue
            all_chunks.extend(file_chunks)
            source_summary.append(
                {"filename": f.filename, "chunks": len(file_chunks)}
            )

        if not all_chunks:
            raise RuntimeError("Extraction produced 0 chunks.")

        texts = [c[0] for c in all_chunks]
        sources = [c[1] for c in all_chunks]

        # 3. Heavy work runs in a worker thread so the event loop stays
        #    responsive for /auth/me, status polls, health checks, etc.
        embeddings = await asyncio.to_thread(_embed_chunks, texts)

        # 4. Build + write index (also blocking — runs in thread).
        await asyncio.to_thread(
            _write_index_files,
            embeddings,
            texts,
            sources,
            settings.faiss_index_path,
        )

        # 5. Hot-reload into the singleton (light, but uses threading lock).
        await asyncio.to_thread(rag_engine.reload)

        # 6. Push to S3 so other App Runner replicas + future restarts pick
        #    up this index. No-op in dev (USE_S3_FOR_FAISS=false).
        from app.services import s3_store

        await asyncio.to_thread(s3_store.push_from_local)

        # 7. Settle build row.
        build = await _get_build(db, build_id)
        if build is not None:
            build.status = IndexBuildStatus.COMPLETE
            build.chunk_count = len(all_chunks)
            build.source_files = source_summary
            build.completed_at = datetime.now(timezone.utc)
            await db.commit()
        logger.info("index_build_complete chunks=%d", len(all_chunks))

    except Exception as exc:  # noqa: BLE001
        logger.exception("index_build_failed")
        build = await _get_build(db, build_id)
        if build is not None:
            build.status = IndexBuildStatus.FAILED
            build.error_message = str(exc)[:2000]
            build.completed_at = datetime.now(timezone.utc)
            await db.commit()


async def _get_build(db: AsyncSession, build_id: uuid.UUID) -> IndexBuild | None:
    result = await db.execute(select(IndexBuild).where(IndexBuild.id == build_id))
    return result.scalar_one_or_none()


# ── Blocking helpers (run inside asyncio.to_thread) ─────────────────────

def _embed_chunks(texts: list[str]) -> np.ndarray:
    """Synchronous: dispatches to configured embedding provider.
    Must be invoked via asyncio.to_thread so the event loop doesn't block."""
    return rag_engine.embed_batch(texts)


def _write_index_files(
    embeddings: np.ndarray,
    texts: list[str],
    sources: list[str],
    idx_dir: Path,
) -> None:
    """Synchronous: build FAISS, write to temp files, swap atomically."""
    import faiss

    dim = int(embeddings.shape[1])
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings)

    idx_dir.mkdir(parents=True, exist_ok=True)
    backups_dir = idx_dir / "backups"
    backups_dir.mkdir(exist_ok=True)

    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    tmp_index = idx_dir / f".tmp_{ts}_{INDEX_FILENAME}"
    tmp_meta = idx_dir / f".tmp_{ts}_{METADATA_FILENAME}"
    faiss.write_index(index, str(tmp_index))
    with tmp_meta.open("wb") as fh:
        pickle.dump({"texts": texts, "sources": sources}, fh)

    live_index = idx_dir / INDEX_FILENAME
    live_meta = idx_dir / METADATA_FILENAME
    if live_index.exists():
        shutil.move(str(live_index), str(backups_dir / f"{ts}_{INDEX_FILENAME}"))
    if live_meta.exists():
        shutil.move(str(live_meta), str(backups_dir / f"{ts}_{METADATA_FILENAME}"))
    tmp_index.rename(live_index)
    tmp_meta.rename(live_meta)
